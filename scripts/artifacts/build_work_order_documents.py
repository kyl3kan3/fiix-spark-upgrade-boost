from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
DOWNLOADS = ROOT / "public" / "templates" / "downloads"
DOCX_OUTPUT = DOWNLOADS / "work-order-template.docx"
PDF_OUTPUT = DOWNLOADS / "work-order-template.pdf"

NAVY = "123A63"
BLUE = "176B87"
PALE_BLUE = "EAF4F7"
PALE_GRAY = "F4F6F8"
MID_GRAY = "667085"
GRID = "CCD5DF"
WHITE = "FFFFFF"
BLACK = "17212B"


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size=5):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end"):
        tag = f"w:{edge}"
        border = tc_borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            tc_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            set_cell_border(cell)


def set_run_font(run, size=10, bold=False, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_doc_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, NAVY, 10, 5),
        ("Heading 3", 11, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_field_cell(cell, label: str, value: str = "", fill: str = WHITE, lines: int = 0):
    cell.text = ""
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(label.upper())
    set_run_font(label_run, size=7.5, bold=True, color=MID_GRAY)
    if value:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        value_run = p2.add_run(value)
        set_run_font(value_run, size=9.5, color=BLACK)
    for _ in range(lines):
        blank = cell.add_paragraph(" ")
        blank.paragraph_format.space_after = Pt(7)


def add_section_heading(doc: Document, number: str, title: str, note: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    badge = p.add_run(f"{number}  ")
    set_run_font(badge, size=9, bold=True, color=BLUE)
    text = p.add_run(title)
    set_run_font(text, size=13, bold=True, color=NAVY)
    if note:
        note_run = p.add_run(f"  {note}")
        set_run_font(note_run, size=8.5, color=MID_GRAY)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = p.add_run("MaintenEase  |  Free maintenance work order template")
    set_run_font(left, size=8, color=MID_GRAY)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.15))
    p.add_run("\t")
    right = p.add_run("maintenease.com/templates/work-order-template")
    set_run_font(right, size=8, color=MID_GRAY)


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("MAINTENANCE OPERATIONS  /  FIELD FORM")
    set_run_font(run, size=8, bold=True, color=MID_GRAY)


def build_docx():
    DOWNLOADS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "Maintenance Work Order Template"
    doc.core_properties.subject = "Editable maintenance work order for request, planning, execution, and close-out"
    doc.core_properties.author = "MaintenEase"
    configure_doc_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header(section)
    add_footer(section)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(1)
    run = kicker.add_run("MAINTENEASE  /  FREE TEMPLATE")
    set_run_font(run, size=8.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Maintenance Work Order")
    set_run_font(run, size=26, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Request  •  Plan  •  Execute  •  Verify  •  Close")
    set_run_font(run, size=10.5, color=MID_GRAY)

    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [2160, 2160, 2160, 2880])
    labels = [
        ("Work order ID", "WO-________________"),
        ("Priority", "[ ] Emergency  [ ] High  [ ] Medium  [ ] Low"),
        ("Status", "____________________"),
        ("Requested date / time", "________________________"),
        ("Site / building", "____________________"),
        ("Area / room", "____________________"),
        ("Asset ID", "____________________"),
        ("Asset name", "________________________"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], labels):
        set_field_cell(cell, label, value, fill=PALE_GRAY)

    add_section_heading(doc, "01", "Request and triage", "Record the original symptom before diagnosis")
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    labels = [
        ("Requester", "____________________________"),
        ("Department / tenant", "____________________________"),
        ("Contact", "____________________________"),
        ("Service impact", "[ ] Safety  [ ] Production  [ ] Occupant  [ ] Compliance  [ ] None"),
        ("Requested completion", "____________________________"),
        ("Request channel", "[ ] Portal  [ ] Phone  [ ] Email  [ ] Inspection  [ ] Other"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], labels):
        set_field_cell(cell, label, value)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_field_cell(table.cell(0, 0), "Problem or work requested", "Describe what was observed, where, when, and the effect on operations.", lines=2)

    add_section_heading(doc, "02", "Planning and approval", "Define safe scope, owner, dates, parts, and permits")
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    labels = [
        ("Assigned technician / vendor", "____________________________"),
        ("Planner / supervisor", "____________________________"),
        ("Target start", "____________________________"),
        ("Due date / service level", "____________________________"),
        ("Estimated labor hours", "____________________________"),
        ("Estimated downtime", "____________________________"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], labels):
        set_field_cell(cell, label, value)

    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_field_cell(table.cell(0, 0), "Safety and access", "[ ] LOTO  [ ] Permit  [ ] PPE  [ ] Barricade  [ ] Escort  [ ] None")
    set_field_cell(table.cell(0, 1), "Required tools and parts", "_______________________________________________")
    set_field_cell(table.cell(1, 0), "Procedure / scope of work", "Reference job plan, manual, drawing, or required steps.", lines=2)
    set_field_cell(table.cell(1, 1), "Approval", "Approved by: ____________________  Date: ______________", lines=2)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Maintenance Work Order  /  Execution and Close-Out")
    set_run_font(run, size=18, bold=True, color=NAVY)

    add_section_heading(doc, "03", "Execution", "Capture actual time, findings, parts, and work performed")
    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    labels = [
        ("Work started", "__________________"),
        ("Work completed", "__________________"),
        ("Asset returned to service", "__________________"),
        ("Actual downtime", "__________________"),
        ("Meter at service", "__________________"),
        ("Failure code", "__________________"),
        ("Cause code", "__________________"),
        ("Resolution code", "__________________"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], labels):
        set_field_cell(cell, label, value, fill=PALE_GRAY)

    table = doc.add_table(rows=3, cols=4)
    set_table_geometry(table, [2700, 2160, 2160, 2340])
    headers = ["Technician / vendor", "Date", "Hours", "Labor cost"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_field_cell(cell, header, fill=PALE_BLUE)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells:
            set_field_cell(cell, "", " ")

    table = doc.add_table(rows=3, cols=5)
    set_table_geometry(table, [1800, 2880, 1440, 1620, 1620])
    headers = ["Part number", "Part / material", "Qty", "Unit cost", "Total"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_field_cell(cell, header, fill=PALE_BLUE)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells:
            set_field_cell(cell, "", " ")

    add_section_heading(doc, "04", "Diagnosis, repair, and verification")
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_field_cell(table.cell(0, 0), "Cause / condition found", "Record evidence, measurements, and the most likely cause.", lines=2)
    set_field_cell(table.cell(0, 1), "Work performed", "Record repair, adjustment, replacement, or inspection completed.", lines=2)
    set_field_cell(table.cell(1, 0), "Test and return-to-service result", "Acceptance criterion: ______________________________", lines=1)
    set_field_cell(table.cell(1, 1), "Follow-up work", "[ ] None  [ ] Corrective WO  [ ] PM change  [ ] Engineering review", lines=1)

    add_section_heading(doc, "05", "Close-out", "Make the completed record useful to the next person")
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    labels = [
        ("Final status", "[ ] Completed  [ ] Closed  [ ] On hold  [ ] Cancelled"),
        ("Total labor", "____________ hours  $____________"),
        ("Total parts / vendor", "$________________________"),
        ("Completed by", "__________________  Date: __________"),
        ("Verified by", "__________________  Date: __________"),
        ("Requester notified", "[ ] Yes  [ ] No  Date: __________"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], labels):
        set_field_cell(cell, label, value, fill=PALE_GRAY)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Tip: keep the requester's original symptom, technician finding, cause, repair, parts, labor, downtime, and verification as separate fields. That structure makes asset history searchable and future analysis reliable.")
    set_run_font(run, size=8.5, color=MID_GRAY)

    doc.save(DOCX_OUTPUT)


class WorkOrderDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str):
        super().__init__(
            filename,
            pagesize=letter,
            leftMargin=0.62 * inch,
            rightMargin=0.62 * inch,
            topMargin=0.65 * inch,
            bottomMargin=0.58 * inch,
            title="Maintenance Work Order Template",
            author="MaintenEase",
        )
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="content")
        self.addPageTemplates(PageTemplate(id="work-order", frames=frame, onPage=self.draw_page))

    def draw_page(self, canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 7)
        canvas.setFillColor(colors.HexColor(f"#{MID_GRAY}"))
        canvas.drawString(self.leftMargin, letter[1] - 0.35 * inch, "MAINTENANCE OPERATIONS  /  FIELD FORM")
        canvas.setFont("Helvetica", 7)
        canvas.drawString(self.leftMargin, 0.32 * inch, "MaintenEase  |  Free maintenance work order template")
        canvas.drawRightString(letter[0] - self.rightMargin, 0.32 * inch, f"Page {doc.page}")
        canvas.restoreState()


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle("Kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7.5, leading=9, textColor=colors.HexColor(f"#{BLUE}"), spaceAfter=2),
        "title": ParagraphStyle("TitleCustom", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=23, leading=26, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_LEFT, spaceAfter=2),
        "subtitle": ParagraphStyle("SubtitleCustom", parent=styles["Normal"], fontName="Helvetica", fontSize=9, leading=11, textColor=colors.HexColor(f"#{MID_GRAY}"), spaceAfter=8),
        "section": ParagraphStyle("Section", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=7, spaceAfter=4, keepWithNext=True),
        "label": ParagraphStyle("Label", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=6.5, leading=8, textColor=colors.HexColor(f"#{MID_GRAY}"), spaceAfter=2),
        "value": ParagraphStyle("Value", parent=styles["Normal"], fontName="Helvetica", fontSize=7.8, leading=10, textColor=colors.HexColor(f"#{BLACK}")),
        "table_header": ParagraphStyle("TableHeader", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7, leading=8, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_CENTER),
        "note": ParagraphStyle("Note", parent=styles["Normal"], fontName="Helvetica", fontSize=7.3, leading=9.2, textColor=colors.HexColor(f"#{MID_GRAY}"), spaceBefore=5),
    }


def pdf_field(styles, label: str, value: str = "", blank_lines: int = 0):
    body = f"<b>{label.upper()}</b>"
    if value:
        body += f"<br/><font color='#{BLACK}'>{value}</font>"
    body += "<br/>" * blank_lines
    return Paragraph(body, ParagraphStyle(
        f"field-{label}-{blank_lines}",
        parent=styles["value"],
        fontSize=7.5,
        leading=9.5,
        textColor=colors.HexColor(f"#{MID_GRAY}"),
    ))


def pdf_table(data, widths, row_heights=None, header_rows=0, fill=None, pad=5):
    table = Table(data, colWidths=widths, rowHeights=row_heights, repeatRows=header_rows)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{GRID}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), pad),
        ("RIGHTPADDING", (0, 0), (-1, -1), pad),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    if fill:
        commands.append(("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{fill}")))
    if header_rows:
        commands.append(("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor(f"#{PALE_BLUE}")))
    table.setStyle(TableStyle(commands))
    return table


def section_heading(styles, number: str, title: str, note: str = ""):
    suffix = f"  <font color='#{MID_GRAY}' size='8'>{note}</font>" if note else ""
    return Paragraph(f"<font color='#{BLUE}'>{number}</font>  {title}{suffix}", styles["section"])


def build_pdf():
    DOWNLOADS.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    doc = WorkOrderDocTemplate(str(PDF_OUTPUT))
    story = [
        Paragraph("MAINTENEASE  /  FREE TEMPLATE", styles["kicker"]),
        Paragraph("Maintenance Work Order", styles["title"]),
        Paragraph("Request  •  Plan  •  Execute  •  Verify  •  Close", styles["subtitle"]),
    ]

    meta = [
        [pdf_field(styles, "Work order ID", "WO-________________"), pdf_field(styles, "Priority", "[ ] Emergency  [ ] High  [ ] Medium  [ ] Low"), pdf_field(styles, "Status", "________________"), pdf_field(styles, "Requested", "________________")],
        [pdf_field(styles, "Site / building", "________________"), pdf_field(styles, "Area / room", "________________"), pdf_field(styles, "Asset ID", "________________"), pdf_field(styles, "Asset name", "________________")],
    ]
    story.append(pdf_table(meta, [1.5 * inch, 2.25 * inch, 1.2 * inch, 2.3 * inch], fill=PALE_GRAY))
    story.append(section_heading(styles, "01", "Request and triage", "Record the original symptom before diagnosis"))
    triage = [
        [pdf_field(styles, "Requester", "________________________"), pdf_field(styles, "Department / tenant", "________________________"), pdf_field(styles, "Contact", "________________________")],
        [pdf_field(styles, "Service impact", "[ ] Safety  [ ] Production  [ ] Occupant  [ ] Compliance"), pdf_field(styles, "Requested completion", "________________________"), pdf_field(styles, "Request channel", "[ ] Portal  [ ] Phone  [ ] Email  [ ] Inspection")],
    ]
    story.append(pdf_table(triage, [2.42 * inch] * 3))
    story.append(pdf_table([[pdf_field(styles, "Problem or work requested", "Describe what was observed, where, when, and the effect on operations.", 3)]], [7.26 * inch], row_heights=[0.78 * inch]))
    story.append(section_heading(styles, "02", "Planning and approval", "Define safe scope, owner, dates, parts, and permits"))
    planning = [
        [pdf_field(styles, "Assigned technician / vendor", "________________________"), pdf_field(styles, "Planner / supervisor", "________________________"), pdf_field(styles, "Target start", "________________________")],
        [pdf_field(styles, "Due date / service level", "________________________"), pdf_field(styles, "Estimated labor", "________ hours"), pdf_field(styles, "Estimated downtime", "________ hours")],
    ]
    story.append(pdf_table(planning, [2.42 * inch] * 3))
    story.append(pdf_table([
        [pdf_field(styles, "Safety and access", "[ ] LOTO  [ ] Permit  [ ] PPE  [ ] Barricade  [ ] Escort  [ ] None"), pdf_field(styles, "Required tools and parts", "________________________________________")],
        [pdf_field(styles, "Procedure / scope of work", "Reference job plan, manual, drawing, or required steps.", 3), pdf_field(styles, "Approval", "Approved by: __________________  Date: __________", 3)],
    ], [3.63 * inch, 3.63 * inch], row_heights=[0.44 * inch, 0.84 * inch]))

    story.append(PageBreak())
    story.extend([
        Paragraph("Maintenance Work Order  /  Execution and Close-Out", ParagraphStyle("Page2Title", parent=styles["title"], fontSize=16, leading=19, spaceAfter=5)),
        section_heading(styles, "03", "Execution", "Capture actual time, findings, parts, and work performed"),
    ])
    execution = [
        [pdf_field(styles, "Work started", "________________"), pdf_field(styles, "Work completed", "________________"), pdf_field(styles, "Returned to service", "________________"), pdf_field(styles, "Actual downtime", "________________")],
        [pdf_field(styles, "Meter at service", "________________"), pdf_field(styles, "Failure code", "________________"), pdf_field(styles, "Cause code", "________________"), pdf_field(styles, "Resolution code", "________________")],
    ]
    story.append(pdf_table(execution, [1.815 * inch] * 4, fill=PALE_GRAY))

    labor_headers = [Paragraph(h, styles["table_header"]) for h in ("Technician / vendor", "Date", "Hours", "Labor cost")]
    labor_rows = [labor_headers] + [["", "", "", ""] for _ in range(3)]
    story.append(Spacer(1, 5))
    story.append(pdf_table(labor_rows, [2.3 * inch, 1.55 * inch, 1.4 * inch, 2.01 * inch], row_heights=[0.28 * inch] + [0.31 * inch] * 3, header_rows=1))

    part_headers = [Paragraph(h, styles["table_header"]) for h in ("Part number", "Part / material", "Qty", "Unit cost", "Total")]
    part_rows = [part_headers] + [["", "", "", "", ""] for _ in range(3)]
    story.append(Spacer(1, 5))
    story.append(pdf_table(part_rows, [1.3 * inch, 2.25 * inch, 0.75 * inch, 1.45 * inch, 1.51 * inch], row_heights=[0.28 * inch] + [0.31 * inch] * 3, header_rows=1))

    story.append(section_heading(styles, "04", "Diagnosis, repair, and verification"))
    story.append(pdf_table([
        [pdf_field(styles, "Cause / condition found", "Record evidence, measurements, and likely cause.", 3), pdf_field(styles, "Work performed", "Record repair, adjustment, replacement, or inspection.", 3)],
        [pdf_field(styles, "Test and return-to-service result", "Acceptance criterion: __________________________", 1), pdf_field(styles, "Follow-up work", "[ ] None  [ ] Corrective WO  [ ] PM change  [ ] Engineering review", 1)],
    ], [3.63 * inch, 3.63 * inch], row_heights=[0.78 * inch, 0.5 * inch]))
    story.append(section_heading(styles, "05", "Close-out", "Make the completed record useful to the next person"))
    closeout = [
        [pdf_field(styles, "Final status", "[ ] Completed  [ ] Closed  [ ] On hold  [ ] Cancelled"), pdf_field(styles, "Total labor", "______ hours  $________"), pdf_field(styles, "Parts / vendor", "$________________")],
        [pdf_field(styles, "Completed by", "________________  Date: ________"), pdf_field(styles, "Verified by", "________________  Date: ________"), pdf_field(styles, "Requester notified", "[ ] Yes  [ ] No  Date: ________")],
    ]
    story.append(pdf_table(closeout, [2.42 * inch] * 3, fill=PALE_GRAY))
    story.append(Paragraph("Tip: keep the original symptom, technician finding, cause, repair, parts, labor, downtime, and verification as separate fields. That structure makes asset history searchable and future analysis reliable.", styles["note"]))

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUTPUT)
    print(PDF_OUTPUT)
